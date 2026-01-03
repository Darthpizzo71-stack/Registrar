"""
Utility functions for meeting management
"""
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from io import BytesIO
from .models import Meeting
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_agenda_pdf(meeting: Meeting) -> BytesIO:
    """
    Generate PDF agenda for a meeting.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#0c4a6e',
        spaceAfter=30,
        alignment=TA_CENTER
    )
    story.append(Paragraph(meeting.title, title_style))
    story.append(Spacer(1, 0.2 * inch))
    
    # Meeting info
    info_style = ParagraphStyle(
        'Info',
        parent=styles['Normal'],
        fontSize=12,
        alignment=TA_CENTER
    )
    story.append(Paragraph(f"Date: {meeting.date}", info_style))
    story.append(Paragraph(f"Time: {meeting.time}", info_style))
    story.append(Paragraph(f"Location: {meeting.location}", info_style))
    story.append(Spacer(1, 0.3 * inch))
    
    # Agenda items
    if meeting.sections.exists():
        for section in meeting.sections.all().order_by('order'):
            section_style = ParagraphStyle(
                'Section',
                parent=styles['Heading2'],
                fontSize=16,
                textColor='#0369a1',
                spaceAfter=12
            )
            story.append(Paragraph(section.title, section_style))
            
            for item in section.items.all().order_by('order'):
                item_style = ParagraphStyle(
                    'Item',
                    parent=styles['Normal'],
                    fontSize=12,
                    leftIndent=0.5 * inch,
                    spaceAfter=6
                )
                item_text = f"<b>{item.number or item.order}. {item.title}</b>"
                story.append(Paragraph(item_text, item_style))
                if item.description:
                    desc_style = ParagraphStyle(
                        'Description',
                        parent=styles['Normal'],
                        fontSize=10,
                        leftIndent=0.7 * inch,
                        spaceAfter=12
                    )
                    story.append(Paragraph(item.description, desc_style))
    else:
        # No sections, list items directly
        for item in meeting.agenda_items.all().order_by('order'):
            item_style = ParagraphStyle(
                'Item',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12
            )
            item_text = f"<b>{item.number or item.order}. {item.title}</b>"
            story.append(Paragraph(item_text, item_style))
            if item.description:
                desc_style = ParagraphStyle(
                    'Description',
                    parent=styles['Normal'],
                    fontSize=10,
                    leftIndent=0.3 * inch,
                    spaceAfter=12
                )
                story.append(Paragraph(item.description, desc_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer


def generate_agenda_docx(meeting: Meeting) -> BytesIO:
    """
    Generate DOCX agenda for a meeting.
    """
    doc = Document()
    
    # Title
    title = doc.add_heading(meeting.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meeting info
    info_para = doc.add_paragraph()
    info_para.add_run(f"Date: {meeting.date}\n").bold = True
    info_para.add_run(f"Time: {meeting.time}\n")
    info_para.add_run(f"Location: {meeting.location}\n")
    info_para.add_run(f"Type: {meeting.get_meeting_type_display()}\n")
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if meeting.description:
        doc.add_paragraph(meeting.description)
    
    doc.add_paragraph()  # Spacer
    
    # Agenda items
    if meeting.sections.exists():
        for section in meeting.sections.all().order_by('order'):
            doc.add_heading(section.title, level=1)
            
            for item in section.items.all().order_by('order'):
                item_para = doc.add_paragraph()
                item_para.add_run(f"{item.number or item.order}. {item.title}").bold = True
                
                if item.description:
                    desc_para = doc.add_paragraph(item.description, style='List Bullet')
                    desc_para.paragraph_format.left_indent = Inches(0.5)
    else:
        # No sections, list items directly
        for item in meeting.agenda_items.all().order_by('order'):
            item_para = doc.add_paragraph()
            item_para.add_run(f"{item.number or item.order}. {item.title}").bold = True
            
            if item.description:
                desc_para = doc.add_paragraph(item.description, style='List Bullet')
                desc_para.paragraph_format.left_indent = Inches(0.3)
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_agenda_packet(meeting: Meeting, format: str = 'pdf', include_attachments: bool = True) -> BytesIO:
    """
    Generate complete agenda packet (PDF or DOCX) with optional attachments.
    
    Args:
        meeting: Meeting instance
        format: 'pdf' or 'docx'
        include_attachments: Whether to include document attachments
    
    Returns:
        BytesIO buffer with the generated packet
    """
    if format.lower() == 'docx':
        return generate_agenda_docx(meeting)
    else:
        return generate_agenda_pdf(meeting)


def generate_meeting_ics(meeting: Meeting, base_url: str = 'https://escribe-backend.onrender.com') -> str:
    """
    Generate ICS (iCalendar) file content for a meeting.
    
    Args:
        meeting: Meeting instance
        base_url: Base URL for the application (for links)
    
    Returns:
        ICS file content as string
    """
    from icalendar import Calendar, Event
    from datetime import datetime, timedelta
    from django.utils import timezone
    
    cal = Calendar()
    cal.add('prodid', '-//Escribe Meeting Management//EN')
    cal.add('version', '2.0')
    cal.add('calscale', 'GREGORIAN')
    cal.add('method', 'PUBLISH')
    
    event = Event()
    event.add('summary', meeting.title)
    event.add('description', meeting.description or f"{meeting.meeting_type.title()} Meeting")
    event.add('location', meeting.location)
    
    # Combine date and time
    meeting_datetime = timezone.make_aware(
        timezone.datetime.combine(meeting.date, meeting.time)
    )
    
    # Set start time
    event.add('dtstart', meeting_datetime)
    
    # Set end time (default to 2 hours, can be customized)
    end_datetime = meeting_datetime + timedelta(hours=2)
    event.add('dtend', end_datetime)
    
    # Add timestamps
    event.add('dtstamp', timezone.now())
    event.add('created', meeting.created_at)
    if meeting.updated_at:
        event.add('last-modified', meeting.updated_at)
    
    # Add UID (unique identifier)
    event.add('uid', f'meeting-{meeting.id}@{base_url.replace("https://", "").replace("http://", "")}')
    
    # Add status
    if meeting.status == 'published':
        event.add('status', 'CONFIRMED')
    elif meeting.status == 'draft':
        event.add('status', 'TENTATIVE')
    elif meeting.status == 'cancelled':
        event.add('status', 'CANCELLED')
    
    # Add URL to meeting
    event.add('url', f'{base_url}/api/meetings/{meeting.id}/')
    
    # Add organizer (created_by)
    if meeting.created_by:
        event.add('organizer', f'MAILTO:{meeting.created_by.email or "noreply@example.com"}')
    
    # Add categories
    event.add('categories', [meeting.meeting_type])
    
    # Add alarm/reminder (24 hours before)
    alarm = Event()
    alarm.add('action', 'DISPLAY')
    alarm.add('description', f'Reminder: {meeting.title}')
    alarm.add('trigger', timedelta(hours=-24))
    event.add_component(alarm)
    
    cal.add_component(event)
    
    return cal.to_ical().decode('utf-8')





